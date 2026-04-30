from __future__ import annotations

from dataclasses import dataclass
from datetime import date, datetime
import subprocess
import sys
import tempfile
from pathlib import Path

try:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
except ImportError:  # pragma: no cover - depends on local Python install
    Workbook = None
    Alignment = Border = Font = PatternFill = Side = None

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError:  # pragma: no cover - depends on local Python install
    Document = None
    WD_ALIGN_PARAGRAPH = Pt = None

from .exchange_rate import ExchangeRateLookupError, ExchangeRateResult, fetch_usd_npr_rate
from .generator import StatementConfig, StatementResult
from .utils import (
    amount_to_words_npr,
    amount_to_words_usd,
    format_amount,
    format_long_date,
    format_slash_date,
    iso_date,
    safe_filename,
    write_json,
)


STATEMENT_EXTENSIONS = {".xlsx", ".xls"}
CERTIFICATE_EXTENSIONS = {".docx", ".doc"}


@dataclass(slots=True)
class TemplateEntry:
    name: str
    path: Path


@dataclass(slots=True)
class TemplateCatalog:
    statement_templates: list[TemplateEntry]
    certificate_templates: list[TemplateEntry]


def scan_template_directory(directory: Path) -> TemplateCatalog:
    statement_templates: list[TemplateEntry] = []
    certificate_templates: list[TemplateEntry] = []
    if not directory.exists():
        return TemplateCatalog(statement_templates, certificate_templates)

    for path in sorted(directory.iterdir(), key=lambda item: item.name.lower()):
        suffix = path.suffix.lower()
        if suffix in STATEMENT_EXTENSIONS:
            statement_templates.append(TemplateEntry(path.stem, path))
        elif suffix in CERTIFICATE_EXTENSIONS:
            certificate_templates.append(TemplateEntry(path.stem, path))
    return TemplateCatalog(statement_templates, certificate_templates)


def resolve_exchange_rate(
    issue_date: date,
    mode: str,
    manual_rate: float | None,
    rate_type: str,
) -> ExchangeRateResult:
    if mode == "manual":
        if not manual_rate or manual_rate <= 0:
            raise ExchangeRateLookupError("Enter a valid manual USD/NPR exchange rate before exporting.")
        return ExchangeRateResult(
            rate=manual_rate,
            rate_type=rate_type.lower(),
            source_date=issue_date,
            source_label="Manual USD/NPR Rate",
        )
    return fetch_usd_npr_rate(issue_date, rate_type=rate_type)


def build_payload(
    config: StatementConfig,
    result: StatementResult,
    exchange_rate: ExchangeRateResult,
) -> dict:
    def num_text(value: float, decimals: int = 2) -> str:
        return f"{float(value):.{decimals}f}"

    total_balance_npr = result.final_balance
    equivalent_usd = round(total_balance_npr / exchange_rate.rate, 2)
    rows = [
        {
            "date": iso_date(row.date),
            "description": row.description,
            "cheque_no": row.cheque_no,
            "debit": num_text(row.debit),
            "credit": num_text(row.credit),
            "balance": num_text(row.balance),
            "category": row.category,
            "is_system": row.is_system,
        }
        for row in result.rows
    ]
    return {
        "account": {
            "bank_name": config.bank_name,
            "branch_name": config.branch_name,
            "customer_name": config.customer_name,
            "customer_address": config.customer_address,
            "account_number": config.account_number,
            "account_type": config.account_type,
            "member_id": config.member_id,
            "currency": config.currency,
            "reference_no": config.reference_no,
            "opening_date_iso": iso_date(config.opening_date) if config.opening_date else "",
            "opening_date_slash": format_slash_date(config.opening_date) if config.opening_date else "",
        },
        "statement": {
            "period_from_iso": iso_date(config.start_date),
            "period_to_iso": iso_date(config.end_date),
            "period_from_slash": format_slash_date(config.start_date),
            "period_to_slash": format_slash_date(config.end_date),
            "period_label_slash": f"{format_slash_date(config.start_date)} to {format_slash_date(config.end_date)}",
            "period_label_iso": f"{iso_date(config.start_date)} to {iso_date(config.end_date)}",
            "issue_date_iso": iso_date(result.issue_date),
            "issue_date_slash": format_slash_date(result.issue_date),
            "issue_date_long": format_long_date(result.issue_date, ordinal=False),
            "issue_date_ordinal": format_long_date(result.issue_date, ordinal=True),
            "as_of_iso": iso_date(result.last_transaction_date),
            "as_of_slash": format_slash_date(result.last_transaction_date),
            "as_of_long": format_long_date(result.last_transaction_date, ordinal=False),
            "as_of_ordinal": format_long_date(result.last_transaction_date, ordinal=True),
            "opening_business_date_iso": iso_date(result.opening_business_date),
            "ending_business_date_iso": iso_date(result.ending_business_date),
        },
        "rates": {
            "interest_rate": num_text(config.interest_rate, 4).rstrip("0").rstrip("."),
            "tax_rate": num_text(config.tax_rate, 4).rstrip("0").rstrip("."),
            "usd_npr": num_text(exchange_rate.rate, 4),
            "usd_npr_text": f"{exchange_rate.rate:,.2f}",
            "rate_type": exchange_rate.rate_type,
            "source_label": exchange_rate.source_label,
            "source_date_iso": iso_date(exchange_rate.source_date),
        },
        "summary": {
            "total_deposits": num_text(result.summary.total_deposits),
            "total_withdrawals": num_text(result.summary.total_withdrawals),
            "total_interest": num_text(result.summary.total_interest),
            "total_tax": num_text(result.summary.total_tax),
            "deposit_count": str(result.summary.deposit_count),
            "withdrawal_count": str(result.summary.withdrawal_count),
            "row_count": str(len(result.rows)),
            "final_balance": num_text(total_balance_npr),
            "final_balance_text": format_amount(total_balance_npr),
        },
        "certificate": {
            "total_balance_npr": num_text(total_balance_npr),
            "total_balance_npr_text": format_amount(total_balance_npr),
            "equivalent_usd": num_text(equivalent_usd),
            "equivalent_usd_text": format_amount(equivalent_usd),
            "balance_words_npr": amount_to_words_npr(total_balance_npr),
            "balance_words_usd": amount_to_words_usd(equivalent_usd),
        },
        "statement_rows": rows,
    }


def default_output_name(kind: str, customer_name: str, template_name: str, issue_date: date, suffix: str) -> str:
    name = safe_filename(customer_name)
    template = safe_filename(template_name)
    return f"{kind}_{name}_{issue_date.isoformat()}_{template}{suffix}"


def _resource_root() -> Path:
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        return Path(sys._MEIPASS)
    return Path(__file__).resolve().parent


def _resource_file(name: str) -> Path:
    root = _resource_root()
    candidates = [
        root / name,
        root / "statement_generator" / name,
        Path(__file__).resolve().parent / name,
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return candidates[0]


def _run_export(mode: str, template_path: Path, output_path: Path, payload: dict) -> None:
    script_path = _resource_file("office_export.ps1")
    creationflags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
    with tempfile.TemporaryDirectory(prefix="statement_generator_") as tmp_dir:
        payload_path = Path(tmp_dir) / "payload.json"
        write_json(payload_path, payload)
        command = [
            "powershell",
            "-NoProfile",
            "-ExecutionPolicy",
            "Bypass",
            "-File",
            str(script_path),
            "-Mode",
            mode,
            "-TemplatePath",
            str(template_path),
            "-OutputPath",
            str(output_path),
            "-PayloadPath",
            str(payload_path),
        ]
        completed = subprocess.run(
            command,
            capture_output=True,
            text=True,
            check=False,
            creationflags=creationflags,
        )
    if completed.returncode != 0:
        message = completed.stderr.strip() or completed.stdout.strip() or "Office export failed."
        raise RuntimeError(message)


def export_statement(template_path: Path, output_path: Path, payload: dict) -> None:
    _run_export("statement", template_path, output_path, payload)


def export_certificate(template_path: Path, output_path: Path, payload: dict) -> None:
    _run_export("certificate", template_path, output_path, payload)


def export_normal_statement(output_path: Path, payload: dict) -> None:
    if Workbook is None:
        raise RuntimeError("Normal Excel export dependency is missing in this environment.")

    wb = Workbook()
    ws = wb.active
    ws.title = "Statement"

    title_fill = PatternFill("solid", fgColor="1F4E78")
    header_fill = PatternFill("solid", fgColor="D9EAF7")
    thin = Side(style="thin", color="9AA5B1")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    ws.merge_cells("A1:F1")
    ws["A1"] = "BANK STATEMENT"
    ws["A1"].font = Font(size=16, bold=True, color="FFFFFF")
    ws["A1"].fill = title_fill
    ws["A1"].alignment = Alignment(horizontal="center")

    ws["A3"] = "Bank"
    ws["B3"] = payload["account"]["bank_name"]
    ws["D3"] = "Branch"
    ws["E3"] = payload["account"]["branch_name"]
    ws["A4"] = "Name"
    ws["B4"] = payload["account"]["customer_name"]
    ws["D4"] = "Account No."
    ws["E4"] = payload["account"]["account_number"]
    ws["A5"] = "Address"
    ws["B5"] = payload["account"]["customer_address"]
    ws["D5"] = "Account Type"
    ws["E5"] = payload["account"]["account_type"]
    ws["A6"] = "Statement Period"
    ws["B6"] = payload["statement"]["period_label_slash"]
    ws["D6"] = "Issue Date"
    ws["E6"] = payload["statement"]["issue_date_slash"]
    ws["A7"] = "Interest Rate"
    ws["B7"] = f"{payload['rates']['interest_rate']}%"
    ws["D7"] = "Tax Rate"
    ws["E7"] = f"{payload['rates']['tax_rate']}%"

    headers = ["Date", "Description", "Cheque No.", "Debit", "Credit", "Balance"]
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=9, column=col, value=header)
        cell.font = Font(bold=True)
        cell.fill = header_fill
        cell.border = border
        cell.alignment = Alignment(horizontal="center")

    row_index = 10
    for item in payload["statement_rows"]:
        date_cell = ws.cell(row=row_index, column=1, value=datetime.strptime(item["date"], "%Y-%m-%d").date())
        date_cell.number_format = "yyyy-mm-dd"
        ws.cell(row=row_index, column=2, value=item["description"])
        cheque_text = str(item["cheque_no"]).strip()
        cheque_value = int(cheque_text) if cheque_text.isdigit() else None
        cheque_cell = ws.cell(row=row_index, column=3, value=cheque_value)
        cheque_cell.number_format = "0"
        debit = float(item["debit"])
        credit = float(item["credit"])
        balance = float(item["balance"])
        ws.cell(row=row_index, column=4, value=debit if debit > 0 else None)
        ws.cell(row=row_index, column=5, value=credit if credit > 0 else None)
        ws.cell(row=row_index, column=6, value=balance)
        for col in range(1, 7):
            cell = ws.cell(row=row_index, column=col)
            cell.border = border
            if col in (4, 5, 6) and cell.value is not None:
                cell.number_format = '#,##0.00'
        row_index += 1

    summary_row = row_index + 1
    ws[f"A{summary_row}"] = "Summary"
    ws[f"A{summary_row}"].font = Font(bold=True)
    ws[f"A{summary_row + 1}"] = "Total Deposits"
    ws[f"B{summary_row + 1}"] = float(payload["summary"]["total_deposits"])
    ws[f"D{summary_row + 1}"] = "Total Withdrawals"
    ws[f"E{summary_row + 1}"] = float(payload["summary"]["total_withdrawals"])
    ws[f"A{summary_row + 2}"] = "Interest"
    ws[f"B{summary_row + 2}"] = float(payload["summary"]["total_interest"])
    ws[f"D{summary_row + 2}"] = "Tax"
    ws[f"E{summary_row + 2}"] = float(payload["summary"]["total_tax"])
    ws[f"A{summary_row + 3}"] = "Final Balance"
    ws[f"B{summary_row + 3}"] = float(payload["summary"]["final_balance"])

    for row in range(summary_row + 1, summary_row + 4):
        for col in (2, 5):
            ws.cell(row=row, column=col).number_format = '#,##0.00'

    widths = {"A": 14, "B": 42, "C": 16, "D": 14, "E": 14, "F": 16}
    for column, width in widths.items():
        ws.column_dimensions[column].width = width
    ws.freeze_panes = "A10"

    wb.save(output_path)


def export_normal_certificate(output_path: Path, payload: dict) -> None:
    if Document is None:
        raise RuntimeError("Normal Word export dependency is missing in this environment.")

    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BALANCE CERTIFICATE")
    run.bold = True
    run.font.size = Pt(16)

    intro = doc.add_paragraph()
    intro.add_run("To whom it may concern\n").bold = True
    intro.add_run(
        f"This is to certify that the balance in the credit of the under mentioned account holder as on "
        f"{payload['statement']['as_of_ordinal']} is as follows."
    )

    lines = [
        f"Name: {payload['account']['customer_name']}",
        f"Address: {payload['account']['customer_address']}",
        f"Account Number: {payload['account']['account_number']}",
        f"Account Type: {payload['account']['account_type']}",
        f"Currency: {payload['account']['currency']}",
        f"Total Balance: NPR {payload['certificate']['total_balance_npr_text']}",
        f"Equivalent to USD: {payload['certificate']['equivalent_usd_text']}",
        f"In Words USD: {payload['certificate']['balance_words_usd']}",
        f"Exchange Rate on Issue Date: 1 USD = NPR {payload['rates']['usd_npr_text']}",
        f"Issue Date: {payload['statement']['issue_date_slash']}",
    ]
    for line in lines:
        paragraph = doc.add_paragraph()
        paragraph.add_run(line)

    closing = doc.add_paragraph()
    closing.add_run(
        "This certificate has been issued at the request of the account holder without obligation on the part of the institution."
    )

    signature = doc.add_paragraph("\n\nAuthorized Signature")
    signature.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(output_path)
